
##############################################################################

import matplotlib.pyplot as plt
import docx
import glob
import re
#import jieba.analyse
import pandas as pd
import numpy as np
from sklearn.cross_validation import train_test_split
from sklearn.linear_model import LinearRegression


#unfiy the format of data of 2017 annual reports
def getAllFile(path):#convert the form 
    doc_list=[]
    for doc in glob.glob("{}/*.docx".format(path)):
        doc=doc.replace("\\" ,"/")
        doc_list.append(doc)
    return doc_list


def cleanData():
    path=".../17年年报/未处理/1/" 
    chapter1_error=[]
    chapter2_error=[]
    doc_list=getAllFile(path)
    stock_num=0
    for doc in doc_list:
    #for doc in temp_list:
        print(stock_num)
        #name=doc.split("/")[8].split(".")[0]
        file=docx.Document(doc)
        passage=""
        for para in file.paragraphs:
            passage=passage+para.text
        try:
            for result in re.findall(r'详见.*?公司业务概要',str(passage)):
                if len(result)<20:
                    passage=passage.replace(result,"")
        except:
            passage=passage
        try:
            for result in re.findall(r'参见.*?公司业务概要',str(passage)):
                if len(result)<20:
                    passage=passage.replace(result,"")
        except:
            passage=passage
        try:
            for result in re.findall(r'详见.*?经营情况讨论与分析',str(passage)):
                if len(result)<20:
                    passage=passage.replace(result,"")
        except:
            passage=passage
        try:
            for result in re.findall(r'参见.*?经营情况讨论与分析',str(passage)):
                if len(result)<20:
                    passage=passage.replace(result,"")
        except:
            passage=passage
        try:
            for result in re.findall(r'详见.*?重要事项',str(passage)):
                if len(result)<20:
                    passage=passage.replace(result,"")
        except:
            passage=passage
        try:
            for result in re.findall(r'参见.*?重要事项',str(passage)):
                if len(result)<20:
                    passage=passage.replace(result,"")
        except:
            passage=passage
        chapter_info1=re.findall(r'公司业务概要.*?经营情况讨论与分析',str(passage),re.DOTALL)      
        if chapter_info1==[]:                                                                    
            chapter1_error.append(doc)      
          #  status1=0 
            chapter1=""                                                       
        else:
            count=[len(info) for info in chapter_info1]                                          
            chapter1=chapter_info1[count.index(max(count))] 
           # status1=1                                     
            if len(chapter1)<195:                                                                
                chapter1=""
                chapter1_error.append(doc)
         #       status1=0
        chapter_info2=re.findall(r'经营情况讨论与分析.*?重要事项',str(passage),re.DOTALL)
        #status2=1
        if chapter_info2==[]:
            chapter2_error.append(doc)
           # status2=0
            chapter2=""
        else:
            count=[len(info) for info in chapter_info2]                                          
            chapter2=chapter_info2[count.index(max(count))]                                      
            if len(chapter2)<300:                                                                
                chapter2=""
                chapter2_error.append(doc)
               # status2=0
        stock_num+=1
        
      
def getMatrix(load_path,saved_path):
    data1=pd.read_csv(".../关键字表/chapter1_keywords.csv",engine='python')#这个engine='python'可以能读取地址有中文字符串的文件
    data2=pd.read_csv(".../关键字表/chapter2_keywords.csv",engine='python')
    del data1['Unnamed: 0']
    del data2['Unnamed: 0']
    chapter1_words=list(data1.word.values)
    chapter2_words=list(data2.word.values)
    total_word=len(chapter1_words)+len(chapter2_words)
    doc_list=getAllFile(load_path)
    x=np.ones((len(doc_list),total_word))    
    stock_num=0
    for doc in doc_list:
        print(stock_num)
        file=docx.Document(doc)
        passage=""
        for para in file.paragraphs:
            passage=passage+para.text
        try:
            for result in re.findall(r'详见.*?公司业务概要',str(passage)):
                if len(result)<20:
                    passage=passage.replace(result,"")
        except:
            passage=passage
        try:
            for result in re.findall(r'参见.*?公司业务概要',str(passage)):
                if len(result)<20:
                    passage=passage.replace(result,"")
        except:
            passage=passage
        try:
            for result in re.findall(r'详见.*?经营情况讨论与分析',str(passage)):
                if len(result)<20:
                    passage=passage.replace(result,"")
        except:
            passage=passage
        try:
            for result in re.findall(r'参见.*?经营情况讨论与分析',str(passage)):
                if len(result)<20:
                    passage=passage.replace(result,"")
        except:
            passage=passage
        try:
            for result in re.findall(r'详见.*?重要事项',str(passage)):
                if len(result)<20:
                    passage=passage.replace(result,"")
        except:
            passage=passage
        try:
            for result in re.findall(r'参见.*?重要事项',str(passage)):
                if len(result)<20:
                    passage=passage.replace(result,"")
        except:
            passage=passage
        chapter_info1=re.findall(r'公司业务概要.*?经营情况讨论与分析',str(passage),re.DOTALL)      
        if chapter_info1==[]:                                                                    
            print("chapter1 error")      
          #  status1=0 
            chapter1=""                                                       
        else:
            count=[len(info) for info in chapter_info1]                                          
            chapter1=chapter_info1[count.index(max(count))] 
           # status1=1                                     
            if len(chapter1)<300:                                                                
                chapter1=""
                print("chapter1 error")
         #       status1=0
        chapter_info2=re.findall(r'经营情况讨论与分析.*?重要事项',str(passage),re.DOTALL)
        #status2=1
        if chapter_info2==[]:
            print("chapter2 error")
           # status2=0
            chapter2=""
        else:
            count=[len(info) for info in chapter_info2]                                          
            chapter2=chapter_info2[count.index(max(count))]                                      
            if len(chapter2)<300:                                                                
                chapter2=""
                print("chapter2 error") 
               # status2=0
        for word_code in range(len(chapter1_words)):
            word_count=chapter1.count(chapter1_words[word_code])
            x[stock_num][word_code]=word_count
        for word_code in range(len(chapter2_words)):
            word_count=chapter2.count(chapter2_words[word_code])
            x[stock_num][len(chapter1_words)+word_code]=word_count    
        stock_num+=1
    np.savetxt(saved_path,x)
   

#fit the test data
#apply the model to predict the EPS in 2018 based on the data from 2017 annual reports
load_path1=".../17年年报/已处理/" 
saved_path1=".../模型预测/x.txt"
getMatrix(load_path1,saved_path1) 
doc_list=getAllFile(load_path1)
stock_earning=pd.read_excel(".../模型预测/A股18年每股净利润.xlsx")
x=np.loadtxt(saved_path1)
y=np.zeros((len(x),1))
count=0
for doc in doc_list:
    try:
        name=doc.split("/")[9].split("17")[0]
        print(name)
        earning=stock_earning[stock_earning['name']==name].EPS.values[0]
        y[count][0]=earning
        count+=1
    except:
        count+=1


k=np.loadtxt(".../模型预测/k.txt")
b=np.loadtxt(".../模型预测/b.txt")
y_pred=np.dot(x, k)+b
stock_score_earning17=pd.DataFrame(columns=('stock','pred_earning','earning'))
for count in range(len(doc_list)):
    name=doc_list[count].split("/")[9].split("16")[0]
    pred_earning=y_pred[count]
    earning=y[count]
    stock_score_earning17=stock_score_earning17.append({"stock":str(name),"pred_earning":str(pred_earning),'earning':earning},ignore_index=True)
stock_score_earning17.to_excel(".../模型预测/A股18年预测净利润排名.xlsx")
